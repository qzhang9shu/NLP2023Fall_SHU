# coding:utf-8
import sys
from PyQt5.QtCore import Qt, QRect, QUrl
from PyQt5.QtGui import QIcon, QPainter, QImage, QBrush, QColor, QFont, QDesktopServices
from PyQt5.QtWidgets import QApplication, QFrame, QStackedWidget, QHBoxLayout, QLabel
from PyQt5.QtCore import QUrl, QSize
from qfluentwidgets import (NavigationAvatarWidget, NavigationItemPosition, MessageBox, FluentWindow,
                            SplashScreen)
from qfluentwidgets import FluentIcon as FIF
from qfluentwidgets import (NavigationInterface,NavigationItemPosition, NavigationWidget, MessageBox,
                            isDarkTheme, setTheme, Theme, qrouter)
from qfluentwidgets import FluentIcon as FIF
from qframelesswindow import FramelessWindow, TitleBar
from use import AttentionLSTM
from PyQt5.Qt import *
from PyQt5.QtWidgets import *
import cv2
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from bert import begin5
import sys
import warnings
from gallery import GalleryInterface
warnings.filterwarnings("ignore")
from PyQt5.QtWidgets import QApplication, QWidget, QAction, QHBoxLayout, QLabel
from qfluentwidgets import (StateToolTip, InfoBadge, ToolTipFilter, PushButton, PixmapLabel,
                            InfoBar, InfoBarIcon, FluentIcon, InfoBarPosition, ProgressBar,
                            IndeterminateProgressBar, SpinBox, ProgressRing, IndeterminateProgressRing)
from qfluentwidgets import RoundMenu, setTheme, Theme, Action, MenuAnimationType, MenuItemDelegate, CheckableMenu, MenuIndicatorType
from qfluentwidgets import FluentIcon as FIF
from file import Ui_File
from YU import Ui_YU
from setting import Ui_Setting
from PyQt5.QtWidgets import QApplication, QMainWindow
from PyQt5.QtWidgets import QApplication, QWidget
from PyQt5.QtGui import QPixmap,QPainter
from PyQt5.QtCore import Qt
from use import begin,begin3
from PyQt5.QtGui import QIcon
from sample_card import signalBus
from PyQt5.QtWidgets import QApplication, QWidget, QHBoxLayout
from sample_card import SampleCardView
from qfluentwidgets import ScrollArea, isDarkTheme, FluentIcon
from qfluentwidgets import PipsScrollButtonDisplayMode, HorizontalPipsPager, VerticalPipsPager, setTheme, Theme,ToolButton
import warnings
warnings.filterwarnings("ignore")
class Setting(QWidget, Ui_Setting):
    def stop(self):
        sys.exit(0)

    def __init__(self):
        super(Setting,self).__init__()
        self.setupUi(self)
        self.view = QWidget(self)
        self.loadSamples()
        self.PushButton_2.clicked.connect(self.stop)
    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk.jpg')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
    def loadSamples(self):
        """ load samples """
        # basic input samples
        basicInputView = SampleCardView(
            self.tr(""), self.view)
        basicInputView.addSampleCard(
            icon='resource/CH.png',
            title="中文输入",
            content=self.tr(
                "对中文语段进行仇恨言论检测"),
            routeKey="widget3",
            index=0
        )
        basicInputView.addSampleCard(
            icon='resource/EN.png',
            title="英文输入",
            content=self.tr(
                "对英文语段进行仇恨言论检测"),
            routeKey="widget1",
            index=1
        )
        self.gridLayout.addWidget(basicInputView,0,0,1,1)
class EmittingStr(QObject):
    textWritten = pyqtSignal(str)  # 定义一个发送str的信号
    def write(self, text):
        self.textWritten.emit(str(text))
class YUA(Ui_YU,GalleryInterface):
    def __init__(self):
        super(YUA,self).__init__()
        self.setupUi(self)
        self.setObjectName('YUA')
        self.PushButton.clicked.connect(self.begin2)
        self.PushButton.clicked.connect(begin)
    def begin2(self):
        sys.stdout = EmittingStr(textWritten=self.outputWritten)
        sys.stderr = EmittingStr(textWritten=self.outputWritten)
        self.textBrowser.clear()
        print('请稍后')
        self.spinner = IndeterminateProgressRing(self)
        self.gridLayout.addWidget(self.spinner, 0, Qt.AlignHCenter)
        loop = QEventLoop(self)
        QTimer.singleShot(500, loop.quit)
        loop.exec()
        self.textBrowser.clear()
        self.spinner.setParent(None)
    def outputWritten(self, text):
        cursor = self.textBrowser.textCursor()
        cursor.movePosition(QTextCursor.End)
        cursor.insertText(text)
        self.textBrowser.setTextCursor(cursor)
        self.textBrowser.ensureCursorVisible()

    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)


from PyQt5 import QtCore, QtGui, QtWidgets
from qfluentwidgets import ImageLabel, PushButton, TextEdit
from YUC import Ui_YUC
class YUC(GalleryInterface, Ui_YUC):
    def __init__(self):
        super(YUC, self).__init__()
        self.setupUi(self)
        self.setObjectName('YUC')
        global i1
        i1 = 0

    def begin2(self):
        from qfluentwidgets import ImageLabel, PushButton, TextEdit
        # self.textBrowser.clear()
        textBrowser = QtWidgets.QTextBrowser(self)
        text1 = ' ' + self.TextEdit.toPlainText()
        textBrowser.setText(text1)
        hight = 1
        global i1
        i1 = i1 + 3
        if len(text1) >= 4:
            textBrowser.setMaximumSize(QtCore.QSize(len(text1) * 24, 50 * hight))
        else:
            textBrowser.setMaximumSize(QtCore.QSize(90, 50 * hight))
        textBrowser.setStyleSheet("    color: black;\n"
                                  "    background-color: rgba(0, 255, 255, 0.8);\n"
                                  "    border: 1px solid rgba(0, 0, 0, 13);\n"
                                  "    border-bottom: 1px solid rgba(0, 0, 0, 100);\n"
                                  "    border-radius: 5px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 24pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser, i1, 1, 1, 1)
        ImageLabel = ImageLabel(self)
        A = QImage('resource/SHU3.png').scaled(40, 50, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        ImageLabel.setImage(A)
        ImageLabel.setMaximumSize(QtCore.QSize(40, 50))
        self.gridLayout.addWidget(ImageLabel, i1, 2, 1, 1)
        spacerItem = QtWidgets.QSpacerItem(50, 50, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem, i1 + 1, 0, 1, 1)
    def begin4(self):
        from qfluentwidgets import ImageLabel, PushButton, TextEdit
        # self.textBrowser.clear()
        textBrowser = QtWidgets.QTextBrowser(self)
        textBrowser2 = QtWidgets.QTextBrowser(self)
        b,a=begin3(self.TextEdit.toPlainText())
        text1 = ' ' + b
        a=' '+a
        textBrowser.setText(text1)
        textBrowser2.setText(a)
        hight = 1
        global i1
        i1 = i1 + 2
        if len(a) >= 4:
            textBrowser2.setMaximumSize(QtCore.QSize(250, 35 * hight))
        else:
            textBrowser2.setMaximumSize(QtCore.QSize(90, 35 * hight))
        textBrowser2.setStyleSheet("    color: red;\n"
                                  "    background-color: rgba(255, 255, 255, 0);\n"
                                  "    border: 0px solid rgba(0, 0, 0, 0);\n"
                                  "    border-bottom: 0px solid rgba(0, 0, 0, 0);\n"
                                  "    border-radius: 0px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 16pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser2, i1-1, 1, 1, 1)
        if len(text1) >= 4:
            textBrowser.setMaximumSize(QtCore.QSize(len(text1) * 24, 50 * hight))
        else:
            textBrowser.setMaximumSize(QtCore.QSize(90, 50 * hight))
        textBrowser.setStyleSheet("    color: black;\n"
                                  "    background-color: rgba(255, 255, 255, 0.8);\n"
                                  "    border: 1px solid rgba(0, 0, 0, 13);\n"
                                  "    border-bottom: 1px solid rgba(0, 0, 0, 100);\n"
                                  "    border-radius: 5px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 24pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser, i1, 1, 1, 1)
        ImageLabel = ImageLabel(self)
        A = QImage('resource/SHU2.png').scaled(40, 50, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        ImageLabel.setImage(A)
        ImageLabel.setMaximumSize(QtCore.QSize(40, 50))
        self.gridLayout.addWidget(ImageLabel, i1, 0, 1, 1)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem, i1 + 1, 0, 1, 1)
    def clear_layout(self):
        while self.gridLayout.count():
            child = self.gridLayout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.gridLayout.addItem(spacerItem, 0, 0, 1, 1)
        self.TextEdit = TextEdit(self)
        self.TextEdit.setObjectName("TextEdit")
        self.gridLayout.addWidget(self.TextEdit, 100, 0, 2, 2)
        self.PushButton = PushButton(self)
        self.PushButton.setObjectName("PushButton")
        self.gridLayout.addWidget(self.PushButton, 100, 2, 1, 1)
        self.PushButton_2 = PushButton(self)
        self.PushButton_2.setObjectName("PushButton_2")
        self.gridLayout.addWidget(self.PushButton_2, 101, 2, 1, 1)
        self.retranslateUi(self)
        global i1
        i1 = 0

    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)

class YUD(GalleryInterface, Ui_YUC):
    def __init__(self):
        super(YUD, self).__init__()
        self.setupUi(self)
        self.setObjectName('YUD')
        global i1
        i1 = 0

    def begin2(self):
        from qfluentwidgets import ImageLabel, PushButton, TextEdit
        self.spinner = IndeterminateProgressRing(self)
        self.gridLayout.addWidget(self.spinner, 99,0,1,1)
        loop = QEventLoop(self)
        QTimer.singleShot(1000, loop.quit)
        loop.exec()
        # self.textBrowser.clear()
        textBrowser = QtWidgets.QTextBrowser(self)
        text1 = ' ' + self.TextEdit.toPlainText()
        textBrowser.setText(text1)
        hight = 1
        global i1
        i1 = i1 + 3
        if len(text1) >= 4:
            textBrowser.setMaximumSize(QtCore.QSize(len(text1) * 40, 50 * hight))
        else:
            textBrowser.setMaximumSize(QtCore.QSize(90, 50 * hight))
        textBrowser.setStyleSheet("    color: black;\n"
                                  "    background-color: rgba(0, 255, 255, 0.8);\n"
                                  "    border: 1px solid rgba(0, 0, 0, 13);\n"
                                  "    border-bottom: 1px solid rgba(0, 0, 0, 100);\n"
                                  "    border-radius: 5px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 24pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser, i1, 1, 1, 1)
        ImageLabel = ImageLabel(self)
        A = QImage('resource/SHU3.png').scaled(40, 50, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        ImageLabel.setImage(A)
        ImageLabel.setMaximumSize(QtCore.QSize(40, 50))
        self.gridLayout.addWidget(ImageLabel, i1, 2, 1, 1)
        spacerItem = QtWidgets.QSpacerItem(50, 50, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem, i1 + 1, 0, 1, 1)
    def begin4(self):
        from qfluentwidgets import ImageLabel, PushButton, TextEdit
        # self.textBrowser.clear()
        textBrowser = QtWidgets.QTextBrowser(self)
        textBrowser2 = QtWidgets.QTextBrowser(self)
        b,a=begin5(self.TextEdit.toPlainText())
        text1 = ' ' + b
        a=' '+a
        textBrowser.setText(text1)
        textBrowser2.setText(a)
        hight = 1
        global i1
        i1 = i1 + 2
        if len(a) >= 4:
            textBrowser2.setMaximumSize(QtCore.QSize(250, 35 * hight))
        else:
            textBrowser2.setMaximumSize(QtCore.QSize(90, 35 * hight))
        textBrowser2.setStyleSheet("    color: red;\n"
                                  "    background-color: rgba(255, 255, 255, 0);\n"
                                  "    border: 0px solid rgba(0, 0, 0, 0);\n"
                                  "    border-bottom: 0px solid rgba(0, 0, 0, 0);\n"
                                  "    border-radius: 0px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 16pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser2, i1-1, 1, 1, 1)
        if len(text1) >= 4:
            textBrowser.setMaximumSize(QtCore.QSize(len(text1) * 40, 50 * hight))
        else:
            textBrowser.setMaximumSize(QtCore.QSize(90, 50 * hight))
        textBrowser.setStyleSheet("    color: black;\n"
                                  "    background-color: rgba(255, 255, 255, 0.8);\n"
                                  "    border: 1px solid rgba(0, 0, 0, 13);\n"
                                  "    border-bottom: 1px solid rgba(0, 0, 0, 100);\n"
                                  "    border-radius: 5px;\n"
                                  "    /* font: 14px \"Segoe UI\", \"Microsoft YaHei\"; */\n"
                                  "    padding: 0px 10px;\n"

                                  'font: 24pt "黑体";'
                                  "    selection-background-color: #00a7b3;")
        self.gridLayout.addWidget(textBrowser, i1, 1, 1, 1)
        ImageLabel = ImageLabel(self)
        A = QImage('resource/SHU2.png').scaled(40, 50, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        ImageLabel.setImage(A)
        ImageLabel.setMaximumSize(QtCore.QSize(40, 50))
        self.gridLayout.addWidget(ImageLabel, i1, 0, 1, 1)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Maximum)
        self.gridLayout.addItem(spacerItem, i1 + 1, 0, 1, 1)
        self.spinner.setParent(None)
    def clear_layout(self):
        while self.gridLayout.count():
            child = self.gridLayout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Minimum, QtWidgets.QSizePolicy.Expanding)
        self.gridLayout.addItem(spacerItem, 0, 0, 1, 1)
        self.TextEdit = TextEdit(self)
        self.TextEdit.setObjectName("TextEdit")
        self.gridLayout.addWidget(self.TextEdit, 100, 0, 2, 2)
        self.PushButton = PushButton(self)
        self.PushButton.setObjectName("PushButton")
        self.gridLayout.addWidget(self.PushButton, 100, 2, 1, 1)
        self.PushButton_2 = PushButton(self)
        self.PushButton_2.setObjectName("PushButton_2")
        self.gridLayout.addWidget(self.PushButton_2, 101, 2, 1, 1)
        self.retranslateUi(self)
        global i1
        i1 = 0

    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
class File(QWidget, Ui_File):
    def __init__(self):
        super(File,self).__init__()
        self.setupUi(self)
        self.PushButton.clicked.connect(self.see)
        self.PushButton_2.clicked.connect(self.open)
    def see(self):
        self.showImage = '1.png'
        self.label.setPixmap(QPixmap(self.showImage).scaled(self.label.width(), self.label.height()))
    def open(self):
        fname, _ = QFileDialog.getOpenFileName(self, '选择图片', '../', 'Image files(*.jpg *.gif *.png*.bmp)')
        self.showImage=fname
        self.label.setPixmap(QPixmap(self.showImage).scaled(self.label.width(), self.label.height()))
    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
from fileB import Ui_fileB
import docx
global doc_content
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return "\n".join(content)
class FileB(GalleryInterface, Ui_fileB):
    def __init__(self):
        super(FileB,self).__init__()
        self.setupUi(self)
        self.setObjectName('FileB')
        self.PushButton.clicked.connect(self.start)
        self.PushButton_2.clicked.connect(self.open1)
        self.PushButton_3.clicked.connect(self.open2)
    def start(self):
        global doc_content
        doc_content,a=begin3(doc_content)
        self.textBrowser.setText(a)
        self.textBrowser.insertPlainText(doc_content)
        self.textBrowser.insertPlainText('\n\n')
    def open1(self):
        fname, _ = QFileDialog.getOpenFileName(self, '选择文本', '../', 'Text files(*.docx)')
        self.textBrowser_2.setText(fname)
        file_path = fname
        global doc_content
        doc_content = read_docx(file_path)
        self.textBrowser.setText(doc_content)
    def open2(self):
        fname2= QFileDialog.getExistingDirectory(self, '选择文本', '../')
        file_path = fname2+'/ENNEW.docx'
        self.textBrowser_3.setText(file_path)
        doc = docx.Document()
        global doc_content
        doc.add_paragraph(doc_content)
        doc.save(file_path)

    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)

class FileC(GalleryInterface, Ui_fileB):
    def __init__(self):
        super(FileC,self).__init__()
        self.setupUi(self)
        self.setObjectName('FileC')
        self.PushButton.clicked.connect(self.start)
        self.PushButton_2.clicked.connect(self.open1)
        self.PushButton_3.clicked.connect(self.open2)
    def start(self):
        global doc_content
        self.spinner = IndeterminateProgressRing(self)
        self.gridLayout.addWidget(self.spinner, 3,3,1,1)
        self.textBrowser.setText('请稍后....')
        loop = QEventLoop(self)
        QTimer.singleShot(1000, loop.quit)
        loop.exec()
        self.spinner.setParent(None)
        doc_content,a=begin5(doc_content)
        self.textBrowser.setText(a)
        self.textBrowser.insertPlainText(doc_content)
        self.textBrowser.insertPlainText('\n\n')
    def open1(self):
        fname, _ = QFileDialog.getOpenFileName(self, '选择文本', '../', 'Text files(*.docx)')
        self.textBrowser_2.setText(fname)
        file_path = fname
        global doc_content
        doc_content = read_docx(file_path)
        self.textBrowser.setText(doc_content)
    def open2(self):
        fname2= QFileDialog.getExistingDirectory(self, '选择文本', '../')
        file_path = fname2+'/新中文.docx'
        self.textBrowser_3.setText(file_path)
        doc = docx.Document()
        global doc_content
        doc.add_paragraph(doc_content)
        doc.save(file_path)

    def paintEvent(self, event):  # set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')  # 换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
class Widget(QWidget):

    def __init__(self, text: str, parent=None):
        super().__init__(parent=parent)
        self.setObjectName(text.replace(' ', '-'))
        self.label = QLabel(text, self)
        self.label.setAlignment(Qt.AlignCenter)
        self.hBoxLayout = QHBoxLayout(self)
        self.hBoxLayout.addWidget(self.label, 1, Qt.AlignCenter)

        # leave some space for title bar
        self.hBoxLayout.setContentsMargins(0, 32, 0, 0)

    def paintEvent(self, event):# set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')#换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)

class Widget2(GalleryInterface):

    def __init__(self, text: str, parent=None):
        super().__init__(parent=parent)
        self.setObjectName('widget2')
        self.view = QWidget(self)
        self.hBoxLayout = QHBoxLayout(self)
        self.loadSamples()
        # leave some space for title bar
        self.hBoxLayout.setContentsMargins(0, 32, 0, 0)
    def loadSamples(self):
        """ load samples """
        # basic input samples
        basicInputView = SampleCardView(
            self.tr(""), self.view)
        basicInputView.addSampleCard(
            icon='resource/CH.png',
            title="中文输入",
            content=self.tr(
                "对中文语段进行仇恨言论检测"),
            routeKey="widget3",
            index=0
        )
        basicInputView.addSampleCard(
            icon='resource/EN.png',
            title="英文输入",
            content=self.tr(
                "对英文语段进行仇恨言论检测"),
            routeKey="widget1",
            index=1
        )
        self.hBoxLayout.addWidget(basicInputView)

    def paintEvent(self, event):# set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')#换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
class Widget1(GalleryInterface):

    def __init__(self, text: str, parent=None):
        super().__init__(parent=parent)
        self.setObjectName('widget1')
        self.view = QWidget(self)
        self.hBoxLayout = QHBoxLayout(self)
        self.loadSamples()
        # leave some space for title bar
        self.hBoxLayout.setContentsMargins(0, 32, 0, 0)
    def loadSamples(self):
        """ load samples """
        # basic input samples
        basicInputView = SampleCardView(
            self.tr(""), self.view)
        basicInputView.addSampleCard(
            icon='resource/MO.png',
            title="模型演示",
            content=self.tr(
                "使用测试集对模型进行评估"),
            routeKey="YUA",
            index=0
        )
        basicInputView.addSampleCard(
            icon='resource/CE.png',
            title="实际检测",
            content=self.tr(
                "输入语句并检测屏蔽"),
            routeKey="YUC",
            index=1
        )
        basicInputView.addSampleCard(
            icon='resource/PI.png',
            title="批量检测",
            content=self.tr(
                "导入文档进行检测"),
            routeKey="FileB",
            index=2
        )
        self.hBoxLayout.addWidget(basicInputView)

    def paintEvent(self, event):# set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')#换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)

class Widget3(GalleryInterface):

    def __init__(self, text: str, parent=None):
        super().__init__(parent=parent)
        self.setObjectName('widget3')
        self.view = QWidget(self)
        self.hBoxLayout = QHBoxLayout(self)
        self.loadSamples()
        # leave some space for title bar
        self.hBoxLayout.setContentsMargins(0, 32, 0, 0)
    def loadSamples(self):
        """ load samples """
        # basic input samples
        basicInputView = SampleCardView(
            self.tr(""), self.view)
        basicInputView.addSampleCard(
            icon='resource/CE.png',
            title="实际检测",
            content=self.tr(
                "输入语句并检测屏蔽"),
            routeKey="YUD",
            index=0
        )
        basicInputView.addSampleCard(
            icon='resource/PI.png',
            title="批量检测",
            content=self.tr(
                "导入文档进行检测"),
            routeKey="FileC",
            index=1
        )
        self.hBoxLayout.addWidget(basicInputView)

    def paintEvent(self, event):# set background_img
        painter = QPainter(self)
        painter.drawRect(self.rect())
        pixmap = QPixmap('resource/bk1.png')#换成自己的图片的相对路径
        painter.drawPixmap(self.rect(), pixmap)
class AvatarWidget(NavigationWidget):
    """ Avatar widget """

    def __init__(self, parent=None):
        super().__init__(isSelectable=False, parent=parent)
        self.avatar = QImage('resource/SHU4.png').scaled(
            24, 24, Qt.KeepAspectRatio, Qt.SmoothTransformation)

    def paintEvent(self, e):
        painter = QPainter(self)
        painter.setRenderHints(
            QPainter.SmoothPixmapTransform | QPainter.Antialiasing)

        painter.setPen(Qt.NoPen)

        if self.isPressed:
            painter.setOpacity(0.7)

        # draw background
        if self.isEnter:
            c = 255 if isDarkTheme() else 0
            painter.setBrush(QColor(c, c, c, 10))
            painter.drawRoundedRect(self.rect(), 5, 5)

        # draw avatar
        painter.setBrush(QBrush(self.avatar))
        painter.translate(8, 6)
        painter.drawEllipse(0, 0, 24, 24)
        painter.translate(-8, -6)

        if not self.isCompacted:
            painter.setPen(Qt.white if isDarkTheme() else Qt.black)
            font = QFont('Segoe UI')
            font.setPixelSize(14)
            painter.setFont(font)
            painter.drawText(QRect(44, 0, 255, 36), Qt.AlignVCenter, '仇恨言论检测')


class CustomTitleBar(TitleBar):
    """ Title bar with icon and title """

    def __init__(self, parent):
        super().__init__(parent)
        # add window icon
        self.iconLabel = QLabel(self)
        self.iconLabel.setFixedSize(18, 18)
        self.hBoxLayout.insertSpacing(0, 10)
        self.hBoxLayout.insertWidget(1, self.iconLabel, 0, Qt.AlignLeft | Qt.AlignBottom)
        self.window().windowIconChanged.connect(self.setIcon)

        # add title label
        self.titleLabel = QLabel(self)
        self.hBoxLayout.insertWidget(2, self.titleLabel, 0, Qt.AlignLeft | Qt.AlignBottom)
        self.titleLabel.setObjectName('titleLabel')
        self.window().windowTitleChanged.connect(self.setTitle)

    def setTitle(self, title):
        self.titleLabel.setText(title)
        self.titleLabel.adjustSize()

    def setIcon(self, icon):
        self.iconLabel.setPixmap(QIcon(icon).pixmap(18, 18))


class Window(FramelessWindow):

    def __init__(self):
        super().__init__()
        self.setTitleBar(CustomTitleBar(self))

        # use dark theme mode
        #setTheme(Theme.DARK)

        self.hBoxLayout = QHBoxLayout(self)
        self.navigationInterface = NavigationInterface(
            self, showMenuButton=True, showReturnButton=True)
        self.stackWidget = QStackedWidget(self)

        # create sub interface
        self.searchInterface = Widget2('Search Interface', self)
        self.picInterface = Widget1('Picture Interface', self)
        self.picInterface1 = YUA()
        self.picInterface2 = YUC()
        self.picInterface3 = FileB()
        self.CNInterface = Widget3('CN Interface', self)
        self.CNInterface2 = YUD()
        self.CNInterface3 = FileC()
        #self.videoInterface = YUA()
        #self.folderInterface = File()
        self.settingInterface = Setting()

        # initialize layout
        self.initLayout()
        self.connectSignalToSlot()
        # add items to navigation interface
        self.initNavigation()
        self.initWindow()
        self.splashScreen.finish()
    def connectSignalToSlot(self):
        signalBus.switchToSampleCard.connect(self.switchToSample)
    def initLayout(self):
        self.hBoxLayout.setSpacing(0)
        self.hBoxLayout.setContentsMargins(0, 0, 0, 0)
        self.hBoxLayout.addWidget(self.navigationInterface)
        self.hBoxLayout.addWidget(self.stackWidget)
        self.hBoxLayout.setStretchFactor(self.stackWidget, 1)

        self.titleBar.raise_()
        self.navigationInterface.displayModeChanged.connect(self.titleBar.raise_)

    def initNavigation(self):
        # enable acrylic effect
        # self.navigationInterface.setAcrylicEnabled(True)

        self.addSubInterface(self.searchInterface, FIF.HOME, '首页')
        self.addSubInterface(self.picInterface, FIF.FONT_SIZE, '英文输入',NavigationItemPosition.SCROLL)
        self.addSubInterface(self.picInterface1, FIF.IOT, '模型演示', parent=self.picInterface)
        self.addSubInterface(self.picInterface2, FIF.CHAT, '实际检测', parent=self.picInterface)
        self.addSubInterface(self.picInterface3, FIF.COPY, '批量检测', parent=self.picInterface)
        #self.addSubInterface(self.videoInterface, FIF.SEARCH, '进行识别')
        self.addSubInterface(self.CNInterface, FIF.LANGUAGE, '中文输入', NavigationItemPosition.SCROLL)
        self.addSubInterface(self.CNInterface2, FIF.CHAT, '实际检测', parent=self.CNInterface)
        self.addSubInterface(self.CNInterface3, FIF.COPY, '批量检测', parent=self.CNInterface)
        self.navigationInterface.addSeparator()

        # add navigation items to scroll area
        #self.addSubInterface(self.folderInterface, FIF.FOLDER, '查看文本', NavigationItemPosition.SCROLL)
        # for i in range(1, 21):
        #     self.navigationInterface.addItem(
        #         f'folder{i}',
        #         FIF.FOLDER,
        #         f'Folder {i}',
        #         lambda: print('Folder clicked'),
        #         position=NavigationItemPosition.SCROLL
        #     )

        # add custom widget to bottom
        self.navigationInterface.addWidget(
            routeKey='avatar',
            widget=AvatarWidget(),
            onClick=self.showMessageBox,
            position=NavigationItemPosition.BOTTOM
        )

        self.addSubInterface(self.settingInterface, FIF.SETTING, '设置', NavigationItemPosition.BOTTOM)

        #!IMPORTANT: don't forget to set the default route key
        qrouter.setDefaultRouteKey(self.stackWidget, self.searchInterface.objectName())

        # set the maximum width
        # self.navigationInterface.setExpandWidth(300)

        self.stackWidget.currentChanged.connect(self.onCurrentInterfaceChanged)
        self.stackWidget.setCurrentIndex(0)

    def initWindow(self):
        self.resize(900, 700)
        self.setWindowIcon(QIcon('resource/SHU2.png'))
        self.setWindowTitle('仇恨言论检测系统')
        self.titleBar.setAttribute(Qt.WA_StyledBackground)
        #self.setMicaEffectEnabled(cfg.get(cfg.micaEnabled))

        # create splash screen
        self.splashScreen = SplashScreen('resource/SHU1.png', self)
        self.splashScreen.setIconSize(QSize(1000, 1000))
        self.splashScreen.raise_()

        desktop = QApplication.desktop().availableGeometry()
        w, h = desktop.width(), desktop.height()
        self.move(w//2 - self.width()//2, h//2 - self.height()//2)
        self.setQss()
        self.show()
        QApplication.processEvents()

    def addSubInterface(self, interface, icon, text: str, position=NavigationItemPosition.TOP, parent=None):
        """ add sub interface """
        self.stackWidget.addWidget(interface)
        self.navigationInterface.addItem(
            routeKey=interface.objectName(),
            icon=icon,
            text=text,
            onClick=lambda: self.switchTo(interface),
            position=position,
            tooltip=text,
            parentRouteKey=parent.objectName() if parent else None
        )
    def setQss(self):
        color = 'dark' if isDarkTheme() else 'light'
        with open(f'resource/{color}/demo.qss', encoding='utf-8') as f:
            self.setStyleSheet(f.read())

    def switchTo(self, widget):
        self.stackWidget.setCurrentWidget(widget)

    def onCurrentInterfaceChanged(self, index):
        widget = self.stackWidget.widget(index)
        self.navigationInterface.setCurrentItem(widget.objectName())
        qrouter.push(self.stackWidget, widget.objectName())

    def showMessageBox(self):
        print('.')

    def resizeEvent(self, e):
        self.titleBar.move(46, 0)
        self.titleBar.resize(self.width()-46, self.titleBar.height())

    def switchToSample(self, routeKey, index):
        """ switch to sample """
        interfaces = self.findChildren(GalleryInterface)
        for w in interfaces:
            if w.objectName() == routeKey:
                self.stackWidget.setCurrentWidget(w)

if __name__ == '__main__':
    QApplication.setHighDpiScaleFactorRoundingPolicy(
        Qt.HighDpiScaleFactorRoundingPolicy.PassThrough)
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)

    app = QApplication(sys.argv)
    app.setAttribute(Qt.AA_DontCreateNativeWidgetSiblings)
    w = Window()
    w.show()
    app.exec_()